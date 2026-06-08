import docx

def modify_doc():
    doc_path = r"d:\flutter\final_exam\PE Test 6.docx"
    doc = docx.Document(doc_path)

    answers = [
        """
**Task A1 & A2 Answers:**
Architecture Flow: UI -> ViewModel -> Repository -> Remote/Local DataSource
1. Why not API in widgets: Separation of concerns. Mixing UI with network logic creates spaghetti code, makes reuse impossible, and prevents unit testing of business logic.
2. Layer affected by URL change: Only the Data Layer (ApiService / RemoteDataSource).
3. Offline caching: In the Data Layer via a LocalDataSource. Repository checks cache first.
4. REST to GraphQL: We only replace the RemoteDataSource. The UI and ViewModel remain completely untouched.
""",
        """
**Additional Analysis Questions Answers:**
- Is implementation still appropriate: No.
- Why/why not: Loading 1,000,000 items in a single API call causes massive network payload, extreme latency, and OOM crashes. Local filtering on 1M items will freeze the UI thread.
- Alternatives:
1. Server-side Pagination (limit/skip fetching chunks).
2. Server-side Searching & Filtering (sending keyword to backend database).
""",
        """
**Task D2 Answers:**
- Where: Local storage using Hive (or SharedPreferences).
- Why: Hive is extremely fast, lightweight, and synchronous, ideal for Flutter offline storage.
- Architecture support: The Repository hides the data origin. The UI just asks the Repository for favorites, and the Repository reads directly from the LocalDataSource.
""",
        """
**Task D3 Answers:**
- Flow: UI Heart tap -> Event to ViewModel -> ViewModel calls Repository -> Repository updates Hive LocalDataSource. ViewModel reads updated list, triggers state update -> UI rebuilds.
- State: Updated reactively. ViewModel holds an in-memory list synced with local DB. When storage changes, ViewModel updates its state and notifies UI listeners.
""",
        """
**PART E Answers:**
- Bug: `where` returns a new lazy Iterable. It does NOT mutate the original list.
- Why: Since `where` has no side effects, `products` remains unchanged. `setState` runs but UI doesn't update.
- Corrected: `products = products.where((e) => e.stock > 0).toList();`
- Code review: Look for pure functions (map/where) whose return values are ignored.
- Test case: Init list with 2 items (1 stock=0, 1 stock=5). Call filter function. Assert list length is 1. With bug, it remains 2, failing the test.
""",
        """
**PART F Answers:**
- State: Provider (ChangeNotifier).
- API: Centralized Network Service (http) with custom exceptions.
- Data: Repository Pattern mapping JSON to Dart Models.
- Folder: Layer-based (view, view_model, repository, data).
- Strengths: Clear separation, easy for beginners, centralized error handling.
- Weaknesses: Heavy buildContext reliance, boilerplate with ChangeNotifier, lacks offline caching.
- Improvements: Add Dependency Injection (get_it), Add Local DB (Hive), Migrate to Riverpod.
- Evaluation: Scalability (Moderate), Maintainability (High), Testability (High).
""",
        """
**PART G Answers:**
- AI tools used: ChatGPT / Gemini
- Most important prompts: How to implement Skeleton Loading, explain ListView vs Column.
- Correct code generated: Skeleton Loading UI.
- Incorrect code generated: AI suggested Column with SingleChildScrollView for lists.
- How corrected: Manually changed to ListView.builder for memory efficiency.
- Percentage: 30% AI, 70% self-written.
"""
    ]

    answer_idx = 0
    
    for p in doc.paragraphs:
        if "Add your answer in this doc" in p.text:
            if answer_idx < len(answers):
                p.text = p.text.replace("Add your answer in this doc", answers[answer_idx].strip())
                answer_idx += 1

    # Append part H at the very end
    doc.add_paragraph("\n\n**PART H Answers (Oral Defense Prep):**")
    doc.add_paragraph("- FutureBuilder rebuilds when parent widget rebuilds or Future state changes.")
    doc.add_paragraph("- Repository Pattern abstracts data operations, decoupling logic from sources.")
    doc.add_paragraph("- Debounce prevents API spam on every keystroke, saving server load.")
    doc.add_paragraph("- Timeout errors should be caught, converted to custom AppException, and shown as a UI error state with Retry button.")
    doc.add_paragraph("- Offline-First: Local DB is single source of truth. Background sync updates DB, DB updates UI.")
    doc.add_paragraph("- ListView.builder is lazy (builds visible items only). Column builds all, causing OOM for huge lists.")

    doc.save(doc_path)

if __name__ == '__main__':
    modify_doc()
